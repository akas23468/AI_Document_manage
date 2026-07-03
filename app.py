"""
OnlyTech Portal - Backend API
Handles: Auth (signup/login), PlantMind AI (Groq), Documents (with real file storage
in Supabase Storage + real PDF/DOCX text extraction), AI-assisted before/after
document updates, Document Approval workflow (now with automatic TITLE-MATCH
routing), Maintenance tickets, Activity log, and dashboard stats.

APPROVAL ROUTING
------------------------------------------------
Every employee document upload (propose_upload) and every proposed document
replacement (propose_document_update) is ALWAYS sent to the admin Approvals
tab for a human decision. PlantMind AI's title-match score and content-quality
score are still computed and attached to the request purely as advisory
context for the admin (shown in the Approvals UI) - they never auto-approve
or auto-reject on their own. Nothing is ever published to the Knowledge Base,
and no file is ever deleted, without an explicit admin decision via
/api/approvals/<id>/decide.

Every submission is written to activity_log so it's fully auditable
from the Overview -> Recent Activity feed, and into document_approvals with
status "Pending" plus the ai_score/ai_reasoning columns if your schema has
them.

Run:
    pip install -r requirements.txt
    cp .env.example .env      # fill in your keys
    python app.py

requirements.txt should include:
    flask
    flask-cors
    python-dotenv
    werkzeug
    supabase
    groq
    pypdf
    python-docx
"""

import io
import os
import re
import json
import uuid
import difflib
import collections
import datetime
from functools import wraps

from flask import Flask, request, jsonify, Response
from flask_cors import CORS
from dotenv import load_dotenv
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename

# ---- PDF / DOCX extraction libraries ---------------------------------------
PYPDF_AVAILABLE = False
DOCX_AVAILABLE = False
try:
    from pypdf import PdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    print("[WARN] pypdf not installed. Run: pip install pypdf")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    print("[WARN] python-docx not installed. Run: pip install python-docx")

# --------------------------------------------------------------------------
# ENV / CONFIG
# --------------------------------------------------------------------------
load_dotenv()

GROQ_API_KEY       = os.getenv("GROQ_API_KEY")
SUPABASE_URL       = os.getenv("SUPABASE_URL")
SUPABASE_KEY       = os.getenv("SUPABASE_KEY")   # service_role key for server-side writes
GROQ_MODEL_NAME     = os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile")
STORAGE_BUCKET      = os.getenv("SUPABASE_STORAGE_BUCKET", "documents")

# Max characters of extracted document text we'll keep / send to Groq per document.
MAX_EXTRACTED_CHARS = 15000

# PlantMind AI scoring thresholds (0-100). These are ONLY used to color-code
# the score badge for the admin in the Approvals tab (green/yellow/red) -
# they no longer drive any auto-approve / auto-reject routing decision.
AI_AUTO_APPROVE_MIN_SCORE = int(os.getenv("AI_AUTO_APPROVE_MIN_SCORE", "90"))
AI_AUTO_REJECT_MAX_SCORE  = int(os.getenv("AI_AUTO_REJECT_MAX_SCORE", "70"))

app = Flask(__name__)
CORS(app)

# ---- Supabase client -------------------------------------------------------
supabase = None
try:
    from supabase import create_client, Client
    if SUPABASE_URL and SUPABASE_KEY:
        supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
    else:
        print("[WARN] SUPABASE_URL / SUPABASE_KEY not set yet.")
except ImportError:
    print("[WARN] supabase package not installed. Run: pip install supabase")

# ---- Groq client -----------------------------------------------------------
groq_client = None
try:
    from groq import Groq
    if GROQ_API_KEY:
        groq_client = Groq(api_key=GROQ_API_KEY)
    else:
        print("[WARN] GROQ_API_KEY not set yet.")
except ImportError:
    print("[WARN] groq package not installed. Run: pip install groq")


def require_supabase(fn):
    @wraps(fn)
    def wrapper(*args, **kwargs):
        if supabase is None:
            return jsonify({"error": "Supabase is not configured. Add SUPABASE_URL and SUPABASE_KEY to .env"}), 500
        return fn(*args, **kwargs)
    return wrapper


def now_iso():
    return datetime.datetime.utcnow().isoformat()


def pick_icon(filename_or_name):
    ext = filename_or_name.rsplit(".", 1)[-1].lower() if "." in filename_or_name else ""
    return {
        "pdf": "📕", "doc": "📝", "docx": "📝", "txt": "📝", "md": "📝",
        "xls": "📊", "xlsx": "📊", "csv": "📊",
        "png": "🖼️", "jpg": "🖼️", "jpeg": "🖼️", "gif": "🖼️", "webp": "🖼️",
        "zip": "🗜️", "json": "🧾", "log": "🧾",
    }.get(ext, "📄")


def log_activity(actor, action, target=None, details=None):
    """Best-effort write to activity_log; never blocks the main request on failure."""
    if supabase is None:
        return
    try:
        supabase.table("activity_log").insert({
            "id": str(uuid.uuid4()),
            "actor": actor,
            "action": action,
            "target": target,
            "details": details,
            "created_at": now_iso(),
        }).execute()
    except Exception as e:
        print("[WARN] activity log failed:", e)


def insert_approval_audit_row(record, ai_result):
    """
    Best-effort write of a Pending approval row into document_approvals,
    including ai_score / ai_reasoning columns if the table has them. Falls
    back to inserting without those two columns if the schema doesn't have
    them yet, so this never breaks the upload flow on an older database.
    """
    if supabase is None:
        return
    enriched = dict(record)
    enriched["ai_score"] = ai_result.get("score")
    enriched["ai_reasoning"] = ai_result.get("reasoning")
    try:
        supabase.table("document_approvals").insert(enriched).execute()
    except Exception:
        try:
            supabase.table("document_approvals").insert(record).execute()
        except Exception as e:
            print("[WARN] Could not write approval audit row:", e)


# --------------------------------------------------------------------------
# TEXT EXTRACTION  ->  PDF / DOCX / plain text
# --------------------------------------------------------------------------
def extract_pdf_text(file_bytes):
    if not PYPDF_AVAILABLE:
        return ""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        parts = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                continue
        return "\n".join(parts).strip()
    except Exception as e:
        print("[WARN] PDF extraction failed:", e)
        return ""


def extract_docx_text(file_bytes):
    if not DOCX_AVAILABLE:
        return ""
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        return "\n".join(parts).strip()
    except Exception as e:
        print("[WARN] DOCX extraction failed:", e)
        return ""


def extract_text(file_bytes, mimetype, filename):
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    mimetype = mimetype or ""

    try:
        if mimetype.startswith("text/") or ext in ("txt", "md", "csv", "json", "log"):
            text = file_bytes.decode("utf-8", errors="ignore")
        elif ext == "pdf" or mimetype == "application/pdf":
            text = extract_pdf_text(file_bytes)
        elif ext == "docx" or mimetype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = extract_docx_text(file_bytes)
        else:
            text = ""
    except Exception as e:
        print("[WARN] extract_text failed:", e)
        text = ""

    if text and len(text) > MAX_EXTRACTED_CHARS:
        text = text[:MAX_EXTRACTED_CHARS]

    return text


def ensure_doc_content(doc):
    existing = (doc.get("content") or "").strip()
    if existing:
        return existing

    file_path = doc.get("file_path")
    if not file_path or supabase is None:
        return ""

    try:
        file_bytes = supabase.storage.from_(STORAGE_BUCKET).download(file_path)
    except Exception as e:
        print(f"[WARN] Could not download {file_path} for extraction:", e)
        return ""

    text = extract_text(
        file_bytes,
        doc.get("file_type"),
        doc.get("file_name") or doc.get("name") or file_path,
    )

    if text:
        try:
            supabase.table("documents").update({"content": text}).eq("id", doc["id"]).execute()
        except Exception as e:
            print("[WARN] Could not cache extracted content:", e)

    return text


def upload_bytes_to_storage(file_bytes, mimetype, original_filename, folder="docs"):
    unique_path = f"{folder}/{uuid.uuid4().hex}_{secure_filename(original_filename)}"
    if supabase:
        supabase.storage.from_(STORAGE_BUCKET).upload(
            unique_path, file_bytes, {"content-type": mimetype or "application/octet-stream"}
        )
    return unique_path


def delete_from_storage(file_path):
    if not file_path or supabase is None:
        return
    try:
        supabase.storage.from_(STORAGE_BUCKET).remove([file_path])
    except Exception as e:
        print(f"[WARN] Could not delete {file_path} from storage:", e)


# --------------------------------------------------------------------------
# TITLE MATCHING  ->  compares a proposed upload's name against the existing
# Knowledge Base ("the manual") - ADVISORY ONLY, shown to the admin as
# context. It no longer decides approve/reject on its own.
# --------------------------------------------------------------------------
def normalize_title(name):
    if not name:
        return ""
    name = name.rsplit(".", 1)[0] if "." in name else name
    name = name.lower()
    name = re.sub(r"[^a-z0-9]+", " ", name).strip()
    return name


def find_best_title_match(candidate_name, existing_docs):
    norm_candidate = normalize_title(candidate_name)
    if not norm_candidate:
        return 0, None

    best_score = 0
    best_doc = None
    for d in existing_docs:
        norm_existing = normalize_title(d.get("name") or "")
        if not norm_existing:
            continue
        ratio = difflib.SequenceMatcher(None, norm_candidate, norm_existing).ratio() * 100
        if ratio > best_score:
            best_score = ratio
            best_doc = d
    return round(best_score), best_doc


def score_upload_by_title_match(filename, doc_name, content_text):
    """
    ADVISORY ONLY: scores how well a proposed employee upload matches an
    existing PUBLISHED document title in the Knowledge Base (the "manual").
    This is shown to the admin as context in the Approvals tab. It never
    decides the outcome - every upload always goes to Pending review.

    Returns: {"decision": "review", "score", "reasoning", "matched_document"}
    """
    existing_docs = []
    if supabase is not None:
        try:
            existing_docs = supabase.table("documents") \
                .select("id,name,version,content,status,file_path,file_type,file_name") \
                .eq("status", "Active").execute().data or []
        except Exception as e:
            print("[WARN] Could not fetch existing documents for title match:", e)

    title_score, matched_doc = find_best_title_match(doc_name or filename, existing_docs)

    # Always "review" - routing decision belongs to the admin, not the score.
    decision = "review"

    if matched_doc:
        reasoning = (f'Title match {title_score}% against existing manual entry '
                     f'"{matched_doc["name"]}" ({matched_doc.get("version", "v1.0")}).')
    else:
        reasoning = "No comparable document title was found in the Knowledge Base manual."

    # Fold in a short content-quality note for the audit trail (non-decisive).
    try:
        content_result = score_upload_with_ai(filename, content_text, doc_name)
        reasoning += f" Content quality read: {content_result.get('score')}/100 - {content_result.get('reasoning')}"
    except Exception as e:
        print("[WARN] Supplementary content scoring failed:", e)

    return {
        "decision": decision,
        "score": title_score,
        "reasoning": reasoning[:900],
        "matched_document": matched_doc,
    }


# --------------------------------------------------------------------------
# AI TEXT ANALYSIS  ->  used by the "Update Docs" before/after dashboard
# --------------------------------------------------------------------------
def analyze_text_with_ai(doc_name, text):
    text = (text or "").strip()
    if not text:
        return {
            "summary": "No readable text could be extracted from this file "
                       "(likely an image, spreadsheet, or scanned document).",
            "key_points": [],
        }

    if groq_client is None:
        return {"summary": "PlantMind AI is not configured on the server.", "key_points": []}

    system_prompt = """You are PlantMind AI, summarizing an internal IT document for a
before/after comparison view an employee and admin will read side by side.
Reply with ONLY a compact JSON object, nothing else - no markdown fences, no commentary:
{"summary": "<2-3 sentence plain-text summary>", "key_points": ["<point 1>", "<point 2>", "<point 3>"]}
Keep key_points to at most 5 short, concrete bullet points (procedures, settings,
version numbers, deadlines, responsibilities, etc.) - skip it (empty list) if the
document is too short or generic to have distinct key points."""

    try:
        completion = groq_client.chat.completions.create(
            model=GROQ_MODEL_NAME,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Document name: {doc_name}\n\nContent:\n{text[:6000]}"},
            ],
            temperature=0.2,
            max_tokens=400,
        )
        raw = (completion.choices[0].message.content or "").strip()
        if raw.startswith("```"):
            raw = raw.strip("`")
            if raw.lower().startswith("json"):
                raw = raw[4:]
        raw = raw.strip()

        parsed = json.loads(raw)
        summary = str(parsed.get("summary", "")).strip() or "No summary available."
        key_points = parsed.get("key_points", [])
        if not isinstance(key_points, list):
            key_points = []
        key_points = [str(p).strip() for p in key_points if str(p).strip()][:5]

        return {"summary": summary, "key_points": key_points}
    except Exception as e:
        print("[WARN] analyze_text_with_ai failed:", e)
        return {
            "summary": f"PlantMind AI could not analyze this file automatically ({str(e)}).",
            "key_points": [],
        }


# --------------------------------------------------------------------------
# AI UPLOAD SCREENING  ->  ADVISORY ONLY general content-quality read, used
# as supplementary context in score_upload_by_title_match and shown to the
# admin for the "Update Docs" replacement flow via propose_document_update.
# It never decides the outcome on its own anymore.
# --------------------------------------------------------------------------
def score_upload_with_ai(filename, content_text, doc_name):
    """
    Ask Groq to evaluate a freshly uploaded/updated file against OnlyTech's
    knowledge-base standards and return a 0-100 readiness score, PURELY as
    advisory context for the admin. The routing decision is always "review".

    Returns: {"decision": "review", "score": int 0-100, "reasoning": str}
    """
    if groq_client is None:
        return {"decision": "review", "score": 0,
                "reasoning": "PlantMind AI is not configured on the server, so this needs a manual look."}

    text_for_ai = (content_text or "").strip()
    if not text_for_ai:
        return {"decision": "review", "score": 0,
                "reasoning": "No readable text could be extracted from this file (likely an image, spreadsheet, or scanned document)."}

    system_prompt = """You are a compliance reviewer for OnlyTech's internal IT knowledge base.
Score how ready a submitted document is to be published as-is, based on:
- Is it a real, coherent IT/operations document (SOP, policy, runbook, guide, report, etc.)?
- Is it free of placeholder, junk, spam, or test content?
- Does it look complete, professional, and safe to publish without edits?
- If it is a REPLACEMENT for an existing document, does it stay on-topic and consistent
  with the kind of document it is replacing (not a completely unrelated file)?

Reply with ONLY a compact JSON object, nothing else - no markdown fences, no commentary:
{"score": <integer 0-100>, "reasoning": "<one or two plain sentences>"}

This score is advisory context for a human admin reviewer, not a final decision."""

    try:
        completion = groq_client.chat.completions.create(
            model=GROQ_MODEL_NAME,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Proposed document name: {doc_name}\nOriginal filename: {filename}\n\nExtracted content:\n{text_for_ai[:6000]}"},
            ],
            temperature=0.1,
            max_tokens=250,
        )
        raw = (completion.choices[0].message.content or "").strip()
        if raw.startswith("```"):
            raw = raw.strip("`")
            if raw.lower().startswith("json"):
                raw = raw[4:]
        raw = raw.strip()

        parsed = json.loads(raw)
        score = int(parsed.get("score", 50))
        score = max(0, min(100, score))
        reasoning = str(parsed.get("reasoning", "")).strip()[:500]

        if not reasoning:
            reasoning = "PlantMind AI evaluated this document automatically."

        # Always "review" - the score is advisory only.
        return {"decision": "review", "score": score, "reasoning": reasoning}
    except Exception as e:
        print("[WARN] AI upload scoring failed:", e)
        return {"decision": "review", "score": 0,
                "reasoning": f"PlantMind AI could not score this file automatically ({str(e)})."}


# --------------------------------------------------------------------------
# HEALTH CHECK
# --------------------------------------------------------------------------
@app.route("/api/health", methods=["GET"])
def health():
    return jsonify({
        "status": "ok",
        "groq_configured": groq_client is not None,
        "groq_model": GROQ_MODEL_NAME,
        "supabase_configured": supabase is not None,
        "storage_bucket": STORAGE_BUCKET,
        "pypdf_available": PYPDF_AVAILABLE,
        "docx_available": DOCX_AVAILABLE,
        "ai_auto_approve_min_score": AI_AUTO_APPROVE_MIN_SCORE,
        "ai_auto_reject_max_score": AI_AUTO_REJECT_MAX_SCORE,
        "all_uploads_require_admin_approval": True,
        "time": now_iso()
    })


# --------------------------------------------------------------------------
# AUTH  ->  table: users
# --------------------------------------------------------------------------
@app.route("/api/auth/register", methods=["POST"])
@require_supabase
def register():
    data = request.get_json(force=True) or {}
    required = ["employee_id", "name", "email", "department", "role", "password"]
    missing = [f for f in required if not data.get(f)]
    if missing:
        return jsonify({"error": f"Missing fields: {', '.join(missing)}"}), 400

    if not data["email"].lower().endswith("@onlytech.com"):
        return jsonify({"error": "Registration is restricted to @onlytech.com emails"}), 400

    existing = supabase.table("users").select("employee_id").eq("employee_id", data["employee_id"]).execute()
    if existing.data:
        return jsonify({"error": "Employee ID already registered"}), 409

    record = {
        "employee_id": data["employee_id"],
        "name": data["name"],
        "email": data["email"],
        "department": data["department"],
        "designation": data.get("designation", data["role"]),
        "role": data["role"],
        "password_hash": generate_password_hash(data["password"]),
        "twofa_enabled": False,
        "created_at": now_iso(),
    }
    supabase.table("users").insert(record).execute()
    record.pop("password_hash")
    return jsonify({"message": "Account created", "user": record}), 201


@app.route("/api/auth/login", methods=["POST"])
@require_supabase
def login():
    data = request.get_json(force=True) or {}
    employee_id = data.get("employee_id")
    password    = data.get("password")
    if not employee_id or not password:
        return jsonify({"error": "employee_id and password are required"}), 400

    if employee_id == "ADMIN_01" and password == "AdminPass123!":
        res = supabase.table("users").select("*").eq("employee_id", "ADMIN_01").execute()
        if not res.data:
            admin_record = {
                "employee_id": "ADMIN_01",
                "name": "Super Admin",
                "email": "admin@onlytech.com",
                "department": "IT Operations",
                "designation": "Systems Architect",
                "role": "Admin",
                "password_hash": generate_password_hash("AdminPass123!"),
                "twofa_enabled": False,
                "created_at": now_iso(),
            }
            supabase.table("users").insert(admin_record).execute()
            admin_record.pop("password_hash", None)
            return jsonify({"message": "Local Admin Session Created & Seeded", "user": admin_record}), 200

    res = supabase.table("users").select("*").eq("employee_id", employee_id).execute()
    if not res.data:
        return jsonify({"error": "Invalid Employee ID or password"}), 401

    user = res.data[0]
    if not check_password_hash(user["password_hash"], password):
        return jsonify({"error": "Invalid Employee ID or password"}), 401

    user.pop("password_hash", None)
    return jsonify({"message": "Login successful", "user": user}), 200


# --------------------------------------------------------------------------
# DOCUMENTS  ->  tables: documents, document_history
# --------------------------------------------------------------------------
@app.route("/api/documents", methods=["GET"])
@require_supabase
def list_documents():
    res = supabase.table("documents").select("*").order("updated_at", desc=True).execute()
    return jsonify(res.data), 200


@app.route("/api/documents/<doc_id>", methods=["GET"])
@require_supabase
def get_document(doc_id):
    doc = supabase.table("documents").select("*").eq("id", doc_id).execute()
    if not doc.data:
        return jsonify({"error": "Document not found"}), 404
    history = supabase.table("document_history").select("*").eq("document_id", doc_id).order("created_at", desc=True).execute()
    result = doc.data[0]
    result["history"] = history.data
    return jsonify(result), 200


@app.route("/api/documents", methods=["POST"])
@require_supabase
def create_document():
    """JSON-based creation. Used by admin tools / scripts, not the employee upload flow."""
    data = request.get_json(force=True) or {}
    if not data.get("name"):
        return jsonify({"error": "Document name is required"}), 400

    doc_id = str(uuid.uuid4())
    record = {
        "id": doc_id,
        "name": data["name"],
        "icon": pick_icon(data["name"]),
        "version": data.get("version", "v1.0"),
        "status": data.get("status", "Active"),
        "summary": data.get("summary", ""),
        "content": data.get("content", ""),
        "uploaded_by": data.get("uploaded_by", "Unknown"),
        "created_at": now_iso(),
        "updated_at": now_iso(),
    }
    supabase.table("documents").insert(record).execute()
    supabase.table("document_history").insert({
        "id": str(uuid.uuid4()),
        "document_id": doc_id,
        "version": record["version"],
        "updated_by": record["uploaded_by"],
        "changes": "Initial upload",
        "content_snapshot": record["content"],
        "created_at": now_iso(),
    }).execute()
    log_activity(record["uploaded_by"], "uploaded_document", record["name"])
    return jsonify(record), 201


@app.route("/api/documents/summarize", methods=["POST"])
def summarize_document_preview():
    """
    PREVIEW-ONLY endpoint for the plain "Upload File" button in Documents.
    Extracts text (plain text, PDF, or DOCX) and asks Groq for a short
    2-3 sentence summary. Nothing is saved to Supabase here.
    """
    if "file" not in request.files:
        return jsonify({"error": "file is required"}), 400
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    file_bytes = file.read()
    mimetype = file.mimetype or "application/octet-stream"
    content_text = extract_text(file_bytes, mimetype, file.filename)

    if not content_text.strip():
        size_kb = len(file_bytes) / 1024
        return jsonify({
            "summary": f"\"{file.filename}\" is a {mimetype or 'binary'} file (~{size_kb:.1f} KB). "
                       f"Automatic text summarization isn't available for this file type, "
                       f"but the file itself will still be submitted for review.",
            "content_preview": "",
            "can_summarize": False
        }), 200

    result = analyze_text_with_ai(file.filename, content_text)
    return jsonify({
        "summary": result["summary"],
        "content_preview": content_text[:500],
        "can_summarize": True
    }), 200


@app.route("/api/documents/analyze-file", methods=["POST"])
def analyze_uploaded_file():
    """
    PREVIEW-ONLY endpoint used by the "Update Docs" dashboard's right-hand
    ("Proposed New Version") panel. Given a replacement file (multipart,
    field 'file', plus optional 'doc_name'), extracts its text and returns
    an AI summary + key points WITHOUT saving anything. The frontend calls
    /api/documents/<id>/propose-update separately once the employee confirms.
    """
    if "file" not in request.files:
        return jsonify({"error": "file is required"}), 400
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    doc_name = request.form.get("doc_name") or file.filename

    file_bytes = file.read()
    mimetype = file.mimetype or "application/octet-stream"
    text = extract_text(file_bytes, mimetype, file.filename)

    result = analyze_text_with_ai(doc_name, text)
    return jsonify(result), 200


@app.route("/api/documents/<doc_id>/analyze", methods=["GET"])
@require_supabase
def analyze_current_document(doc_id):
    doc = supabase.table("documents").select("*").eq("id", doc_id).execute()
    if not doc.data:
        return jsonify({"error": "Document not found"}), 404

    d = doc.data[0]
    text = ensure_doc_content(d)
    result = analyze_text_with_ai(d.get("name", "Document"), text or d.get("summary", ""))
    return jsonify(result), 200


@app.route("/api/documents/propose-upload", methods=["POST"])
@require_supabase
def propose_upload():
    """
    Employee-facing NEW document upload endpoint (plain Documents tab).

    ALWAYS routes to the admin Approvals tab as a Pending request. Nothing
    is ever published to the Knowledge Base and no uploaded file is ever
    deleted without an explicit admin decision via /api/approvals/<id>/decide.
    PlantMind AI's title-match score is computed and attached purely as
    advisory context for the admin.
    """
    if "file" not in request.files:
        return jsonify({"error": "file is required"}), 400
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    uploaded_by = request.form.get("uploaded_by", "Unknown")
    name = request.form.get("name") or file.filename
    summary = request.form.get("summary", "")

    file_bytes = file.read()
    mimetype = file.mimetype or "application/octet-stream"
    content_text = extract_text(file_bytes, mimetype, file.filename)
    unique_path = upload_bytes_to_storage(file_bytes, mimetype, file.filename, folder="approvals")

    try:
        ai_result = score_upload_by_title_match(file.filename, name, content_text or summary)
    except Exception as e:
        ai_result = {"decision": "review", "score": 0,
                     "reasoning": f"Title-match scoring failed: {e}", "matched_document": None}

    matched_doc = ai_result.get("matched_document")

    # Compute what version this upload would become, and against which
    # document_id, based on whether we found a matching manual entry.
    # (Purely informational until the admin approves.)
    if matched_doc:
        old_version = matched_doc.get("version", "v1.0")
        try:
            num = float(old_version.replace("v", ""))
            proposed_version = f"v{round(num + 0.1, 1)}"
        except Exception:
            proposed_version = "v2.0"
        old_content = ensure_doc_content(matched_doc)
        document_id = matched_doc["id"]
    else:
        old_version = "v0.0"
        proposed_version = "v1.0"
        old_content = ""
        document_id = None

    pending_record = {
        "id": str(uuid.uuid4()),
        "document_id": document_id,
        "doc_name": name,
        "old_version": old_version,
        "proposed_version": proposed_version,
        "old_content": old_content,
        "new_content": content_text,
        "file_path": unique_path,
        "file_name": file.filename,
        "file_type": mimetype,
        "file_size": len(file_bytes),
        "submitted_by": uploaded_by,
        "status": "Pending",
        "created_at": now_iso(),
    }
    insert_approval_audit_row(pending_record, ai_result)

    log_activity(
        uploaded_by, "proposed_edit", name,
        f"Submitted for admin review — title match {ai_result['score']}%: {ai_result['reasoning']}"
    )
    return jsonify({
        "outcome": "review",
        "message": f"\"{name}\" was submitted and is now pending admin approval (title match {ai_result['score']}%, shown to the admin for context).",
        "ai": ai_result,
    }), 202


@app.route("/api/documents/<doc_id>/propose-update", methods=["POST"])
@require_supabase
def propose_document_update(doc_id):
    """
    Employee-facing REPLACEMENT upload for an EXISTING document, from the
    "Update Docs" dashboard.

    ALWAYS routes to the admin Approvals tab as a Pending request. Nothing
    is ever published to the Knowledge Base and no uploaded file is ever
    deleted without an explicit admin decision via /api/approvals/<id>/decide.
    PlantMind AI's content-quality score is computed and attached purely as
    advisory context for the admin.
    """
    existing = supabase.table("documents").select("*").eq("id", doc_id).execute()
    if not existing.data:
        return jsonify({"error": "Document not found"}), 404
    current = existing.data[0]

    if "file" not in request.files:
        return jsonify({"error": "file is required"}), 400
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    submitted_by = request.form.get("submitted_by", "Unknown")
    reason = request.form.get("reason", "")

    file_bytes = file.read()
    mimetype = file.mimetype or "application/octet-stream"
    new_content_text = extract_text(file_bytes, mimetype, file.filename)
    unique_path = upload_bytes_to_storage(file_bytes, mimetype, file.filename, folder="approvals")

    old_version = current.get("version", "v1.0")
    try:
        num = float(old_version.replace("v", ""))
        proposed_version = f"v{round(num + 0.1, 1)}"
    except Exception:
        proposed_version = "v2.0"

    old_content = ensure_doc_content(current)

    try:
        ai_result = score_upload_with_ai(file.filename, new_content_text, current["name"])
    except Exception as e:
        ai_result = {"decision": "review", "score": 0, "reasoning": f"AI scoring failed: {e}"}

    pending_record = {
        "id": str(uuid.uuid4()),
        "document_id": doc_id,
        "doc_name": current["name"],
        "old_version": old_version,
        "proposed_version": proposed_version,
        "old_content": old_content,
        "new_content": new_content_text,
        "file_path": unique_path,
        "file_name": file.filename,
        "file_type": mimetype,
        "file_size": len(file_bytes),
        "submitted_by": submitted_by,
        "status": "Pending",
        "created_at": now_iso(),
    }
    insert_approval_audit_row(pending_record, ai_result)

    log_activity(submitted_by, "proposed_edit", current["name"],
                 f"Needs admin review — PlantMind AI score {ai_result['score']}/100: {reason or ai_result['reasoning']}")
    return jsonify({
        "message": f"Update submitted — PlantMind AI scored this {ai_result['score']}/100 (shown to the admin for context). An admin will review it before it's published.",
        "outcome": "review", "approval_id": pending_record["id"], "proposed_version": proposed_version, "ai": ai_result,
    }), 201


@app.route("/api/documents/upload", methods=["POST"])
@require_supabase
def upload_document():
    """
    Direct admin upload -> stored in Supabase Storage, text extracted, published
    immediately with no approval step. Used by the admin Command Center's
    Knowledge Base "Upload File" button ONLY (admin is already the approver).
    Employee uploads go through /api/documents/propose-upload or
    /api/documents/<id>/propose-update instead, and ALWAYS require an
    explicit admin decision in the Approvals tab.
    """
    if "file" not in request.files:
        return jsonify({"error": "file is required"}), 400
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    uploaded_by = request.form.get("uploaded_by", "Super Admin")
    name = request.form.get("name") or file.filename
    summary = request.form.get("summary", "")

    file_bytes = file.read()
    mimetype = file.mimetype or "application/octet-stream"
    content_text = extract_text(file_bytes, mimetype, file.filename)
    unique_path = upload_bytes_to_storage(file_bytes, mimetype, file.filename)

    doc_id = str(uuid.uuid4())
    record = {
        "id": doc_id,
        "name": name,
        "icon": pick_icon(file.filename),
        "version": "v1.0",
        "status": "Active",
        "summary": summary,
        "content": content_text,
        "file_path": unique_path,
        "file_name": file.filename,
        "file_type": mimetype,
        "file_size": len(file_bytes),
        "uploaded_by": uploaded_by,
        "created_at": now_iso(),
        "updated_at": now_iso(),
    }
    supabase.table("documents").insert(record).execute()
    supabase.table("document_history").insert({
        "id": str(uuid.uuid4()), "document_id": doc_id, "version": "v1.0",
        "updated_by": uploaded_by, "changes": "Initial upload",
        "content_snapshot": content_text, "created_at": now_iso(),
    }).execute()
    log_activity(uploaded_by, "uploaded_document", name, f"Uploaded {file.filename}")
    return jsonify(record), 201


@app.route("/api/documents/<doc_id>/download", methods=["GET"])
@require_supabase
def download_document(doc_id):
    doc = supabase.table("documents").select("*").eq("id", doc_id).execute()
    if not doc.data:
        return jsonify({"error": "Document not found"}), 404
    d = doc.data[0]
    if not d.get("file_path"):
        return jsonify({"error": "No file stored for this document"}), 404
    file_bytes = supabase.storage.from_(STORAGE_BUCKET).download(d["file_path"])
    return Response(
        file_bytes,
        mimetype=d.get("file_type") or "application/octet-stream",
        headers={"Content-Disposition": f'inline; filename="{d.get("file_name") or d["name"]}"'}
    )


@app.route("/api/documents/<doc_id>", methods=["DELETE"])
@require_supabase
def delete_document(doc_id):
    doc = supabase.table("documents").select("*").eq("id", doc_id).execute()
    if not doc.data:
        return jsonify({"error": "Document not found"}), 404
    d = doc.data[0]

    delete_from_storage(d.get("file_path"))

    try:
        supabase.table("document_history").delete().eq("document_id", doc_id).execute()
    except Exception as e:
        print("[WARN] Could not delete document history:", e)

    supabase.table("documents").delete().eq("id", doc_id).execute()

    deleted_by = request.args.get("deleted_by")
    if not deleted_by:
        deleted_by = (request.get_json(silent=True) or {}).get("deleted_by", "Super Admin")
    log_activity(deleted_by, "deleted_document", d["name"], "Document permanently removed from knowledge base")

    return jsonify({"message": "Document deleted", "id": doc_id}), 200


@app.route("/api/documents/<doc_id>/version", methods=["POST"])
@require_supabase
def update_document_version(doc_id):
    """Direct admin version bump."""
    data = request.get_json(force=True) or {}
    doc = supabase.table("documents").select("*").eq("id", doc_id).execute()
    if not doc.data:
        return jsonify({"error": "Document not found"}), 404

    current = doc.data[0]
    try:
        num = float(current["version"].replace("v", ""))
        new_version = f"v{round(num + 0.1, 1)}"
    except Exception:
        new_version = "v2.0"

    updates = {"version": new_version, "status": "Updated", "updated_at": now_iso()}
    if data.get("content"):
        updates["content"] = data["content"]

    supabase.table("documents").update(updates).eq("id", doc_id).execute()
    supabase.table("document_history").insert({
        "id": str(uuid.uuid4()), "document_id": doc_id, "version": new_version,
        "updated_by": data.get("updated_by", "Unknown"),
        "changes": data.get("changes", "Updated document"),
        "content_snapshot": data.get("content", ""), "created_at": now_iso(),
    }).execute()
    log_activity(data.get("updated_by", "Unknown"), "uploaded_document", current["name"], "Version updated")
    return jsonify({"message": "Version updated", "version": new_version}), 200


@app.route("/api/documents/backfill-content", methods=["POST"])
@require_supabase
def backfill_document_content():
    docs = supabase.table("documents").select("*").execute().data or []
    updated, skipped, failed = [], [], []

    for doc in docs:
        if (doc.get("content") or "").strip():
            skipped.append(doc["name"])
            continue
        if not doc.get("file_path"):
            skipped.append(doc["name"])
            continue
        text = ensure_doc_content(doc)
        if text:
            updated.append(doc["name"])
        else:
            failed.append(doc["name"])

    return jsonify({
        "updated": updated,
        "skipped_already_had_content_or_no_file": skipped,
        "failed_to_extract": failed,
    }), 200


# --------------------------------------------------------------------------
# DOCUMENT APPROVALS  ->  table: document_approvals
# --------------------------------------------------------------------------
@app.route("/api/documents/propose", methods=["POST"])
@require_supabase
def propose_document():
    """Propose an edit to an existing document (multipart/form-data)."""
    submitted_by = request.form.get("submitted_by", "Unknown")
    doc_id = request.form.get("document_id") or None
    doc_name = (request.form.get("name") or "").strip()
    new_content_text = request.form.get("content", "")

    file = request.files.get("file")
    file_path = file_name = file_type = None
    file_size = None
    if file and file.filename:
        file_bytes = file.read()
        file_type = file.mimetype or "application/octet-stream"
        file_name = file.filename
        file_size = len(file_bytes)
        file_path = upload_bytes_to_storage(file_bytes, file_type, file.filename, folder="approvals")
        if not new_content_text:
            new_content_text = extract_text(file_bytes, file_type, file.filename)

    old_content = ""
    old_version = "v0.0"
    if doc_id:
        existing = supabase.table("documents").select("*").eq("id", doc_id).execute()
        if not existing.data:
            return jsonify({"error": "Document not found"}), 404
        old_content = existing.data[0].get("content", "") or ""
        old_version = existing.data[0].get("version", "v1.0")
        doc_name = doc_name or existing.data[0]["name"]
        try:
            num = float(old_version.replace("v", ""))
            proposed_version = f"v{round(num + 0.1, 1)}"
        except Exception:
            proposed_version = "v2.0"
    else:
        proposed_version = "v1.0"

    if not doc_name:
        return jsonify({"error": "Document name is required"}), 400

    approval_id = str(uuid.uuid4())
    record = {
        "id": approval_id,
        "document_id": doc_id,
        "doc_name": doc_name,
        "old_version": old_version,
        "proposed_version": proposed_version,
        "old_content": old_content,
        "new_content": new_content_text,
        "file_path": file_path,
        "file_name": file_name,
        "file_type": file_type,
        "file_size": file_size,
        "submitted_by": submitted_by,
        "status": "Pending",
        "created_at": now_iso(),
    }
    supabase.table("document_approvals").insert(record).execute()
    log_activity(submitted_by, "proposed_edit", doc_name, f"Proposed {proposed_version}")
    return jsonify(record), 201


@app.route("/api/approvals", methods=["POST"])
@require_supabase
def propose_edit_json():
    """
    JSON-based edit proposal, used by the employee portal's "Propose New
    Version" dialog on the Documents tab (posts JSON, not multipart).
    """
    data = request.get_json(force=True) or {}
    doc_id = data.get("doc_id")
    doc_name = (data.get("doc_name") or "").strip()
    if not doc_id or not doc_name:
        return jsonify({"error": "doc_id and doc_name are required"}), 400

    record = {
        "id": str(uuid.uuid4()),
        "document_id": doc_id,
        "doc_name": doc_name,
        "old_version": data.get("old_version", "v1.0"),
        "proposed_version": data.get("proposed_version", "v1.1"),
        "old_content": data.get("old_content", ""),
        "new_content": data.get("new_content", ""),
        "file_path": None,
        "file_name": data.get("file_name"),
        "file_size": data.get("file_size"),
        "file_type": None,
        "submitted_by": data.get("submitted_by", "Unknown"),
        "status": "Pending",
        "created_at": now_iso(),
    }
    supabase.table("document_approvals").insert(record).execute()
    log_activity(record["submitted_by"], "proposed_edit", doc_name, data.get("changes", "Proposed an edit"))
    return jsonify(record), 201


@app.route("/api/approvals", methods=["GET"])
@require_supabase
def list_approvals():
    status = request.args.get("status")
    submitted_by = request.args.get("submitted_by")

    q = supabase.table("document_approvals").select("*")
    if status:
        q = q.eq("status", status)
    if submitted_by:
        q = q.eq("submitted_by", submitted_by)
    res = q.order("created_at", desc=True).execute()
    return jsonify(res.data), 200


@app.route("/api/approvals/<approval_id>", methods=["GET"])
@require_supabase
def get_approval(approval_id):
    res = supabase.table("document_approvals").select("*").eq("id", approval_id).execute()
    if not res.data:
        return jsonify({"error": "Approval not found"}), 404
    return jsonify(res.data[0]), 200


@app.route("/api/approvals/<approval_id>/file", methods=["GET"])
@require_supabase
def download_approval_file(approval_id):
    res = supabase.table("document_approvals").select("*").eq("id", approval_id).execute()
    if not res.data:
        return jsonify({"error": "Not found"}), 404
    a = res.data[0]
    if not a.get("file_path"):
        return jsonify({"error": "No file attached"}), 404
    file_bytes = supabase.storage.from_(STORAGE_BUCKET).download(a["file_path"])
    return Response(
        file_bytes,
        mimetype=a.get("file_type") or "application/octet-stream",
        headers={"Content-Disposition": f'inline; filename="{a.get("file_name") or "file"}"'}
    )


@app.route("/api/approvals/<approval_id>/decide", methods=["POST"])
@require_supabase
def decide_approval(approval_id):
    """
    THE ONLY PLACE a proposed upload/update is ever published or discarded.
    Every employee submission (new upload or replacement) lands here as
    Pending and stays that way until an admin explicitly approves or
    rejects it. Approving publishes the file to the Knowledge Base
    (updating the matched document if document_id is set, otherwise
    creating a brand new document); rejecting deletes the uploaded file and
    removes the request.
    """
    data = request.get_json(force=True) or {}
    approve = bool(data.get("approve"))
    decided_by = data.get("decided_by", "Super Admin")

    res = supabase.table("document_approvals").select("*").eq("id", approval_id).execute()
    if not res.data:
        return jsonify({"error": "Approval not found"}), 404
    app_row = res.data[0]

    if app_row["status"] != "Pending":
        return jsonify({"error": "This request was already decided"}), 409

    if approve:
        if app_row["document_id"]:
            doc_id = app_row["document_id"]
            updates = {
                "version": app_row["proposed_version"],
                "status": "Active",
                "content": app_row["new_content"],
                "updated_at": now_iso(),
            }
            if app_row.get("file_path"):
                updates.update({
                    "file_path": app_row["file_path"],
                    "file_name": app_row["file_name"],
                    "file_type": app_row["file_type"],
                    "file_size": app_row["file_size"],
                })
            supabase.table("documents").update(updates).eq("id", doc_id).execute()
        else:
            doc_id = str(uuid.uuid4())
            supabase.table("documents").insert({
                "id": doc_id,
                "name": app_row["doc_name"],
                "icon": pick_icon(app_row.get("file_name") or app_row["doc_name"]),
                "version": app_row["proposed_version"],
                "status": "Active",
                "summary": "",
                "content": app_row["new_content"],
                "file_path": app_row.get("file_path"),
                "file_name": app_row.get("file_name"),
                "file_type": app_row.get("file_type"),
                "file_size": app_row.get("file_size"),
                "uploaded_by": app_row["submitted_by"],
                "created_at": now_iso(),
                "updated_at": now_iso(),
            }).execute()

        supabase.table("document_history").insert({
            "id": str(uuid.uuid4()), "document_id": doc_id, "version": app_row["proposed_version"],
            "updated_by": app_row["submitted_by"], "changes": "Approved edit request.",
            "content_snapshot": app_row["new_content"], "created_at": now_iso(),
        }).execute()

        supabase.table("document_approvals").update({
            "status": "Approved", "decided_at": now_iso(), "decided_by": decided_by,
        }).eq("id", approval_id).execute()

        log_activity(decided_by, "approved_edit", app_row["doc_name"], f"Published {app_row['proposed_version']}")
    else:
        if app_row.get("file_path"):
            delete_from_storage(app_row["file_path"])

        supabase.table("document_approvals").delete().eq("id", approval_id).execute()
        log_activity(decided_by, "rejected_edit", app_row["doc_name"], "Edit request rejected and file removed")

    return jsonify({"message": "Decision recorded", "approve": approve}), 200


# --------------------------------------------------------------------------
# MAINTENANCE TICKETS  ->  table: maintenance_tickets
# --------------------------------------------------------------------------
@app.route("/api/maintenance", methods=["GET"])
@require_supabase
def list_tickets():
    res = supabase.table("maintenance_tickets").select("*").order("created_at", desc=True).execute()
    return jsonify(res.data), 200


@app.route("/api/maintenance", methods=["POST"])
@require_supabase
def create_ticket():
    is_multipart = bool(request.content_type) and "multipart/form-data" in request.content_type
    data = request.form if is_multipart else (request.get_json(force=True) or {})

    required = ["system_name", "issue_type", "description", "severity", "reported_by"]
    missing = [f for f in required if not data.get(f)]
    if missing:
        return jsonify({"error": f"Missing fields: {', '.join(missing)}"}), 400

    file_path = file_name = file_type = None
    file_size = None
    if is_multipart:
        file = request.files.get("file")
        if file and file.filename:
            file_bytes = file.read()
            file_type = file.mimetype or "application/octet-stream"
            file_name = file.filename
            file_size = len(file_bytes)
            file_path = upload_bytes_to_storage(file_bytes, file_type, file.filename, folder="tickets")

    count_res = supabase.table("maintenance_tickets").select("id", count="exact").execute()
    next_num = (count_res.count or 0) + 1042

    record = {
        "id": str(uuid.uuid4()),
        "ticket_id": f"INC-{next_num}",
        "system_name": data["system_name"],
        "issue_type": data["issue_type"],
        "severity": data["severity"],
        "description": data["description"],
        "reported_by": data["reported_by"],
        "status": "Open",
        "resolution_notes": None,
        "file_path": file_path,
        "file_name": file_name,
        "file_type": file_type,
        "file_size": file_size,
        "created_at": now_iso(),
        "resolved_at": None,
    }
    supabase.table("maintenance_tickets").insert(record).execute()
    log_activity(data["reported_by"], "reported_incident", record["ticket_id"], data["system_name"])
    return jsonify(record), 201


@app.route("/api/maintenance/<ticket_id>/file", methods=["GET"])
@require_supabase
def download_ticket_file(ticket_id):
    ticket = supabase.table("maintenance_tickets").select("*").eq("id", ticket_id).execute()
    if not ticket.data:
        return jsonify({"error": "Ticket not found"}), 404
    t = ticket.data[0]
    if not t.get("file_path"):
        return jsonify({"error": "No file attached to this ticket"}), 404
    file_bytes = supabase.storage.from_(STORAGE_BUCKET).download(t["file_path"])
    return Response(
        file_bytes,
        mimetype=t.get("file_type") or "application/octet-stream",
        headers={"Content-Disposition": f'inline; filename="{t.get("file_name") or "attachment"}"'}
    )


@app.route("/api/maintenance/<ticket_id>/resolve", methods=["POST"])
@require_supabase
def resolve_ticket(ticket_id):
    data = request.get_json(force=True) or {}
    notes = data.get("resolution_notes")
    if not notes:
        return jsonify({"error": "resolution_notes is required"}), 400

    ticket = supabase.table("maintenance_tickets").select("*").eq("id", ticket_id).execute()
    if not ticket.data:
        return jsonify({"error": "Ticket not found"}), 404

    supabase.table("maintenance_tickets").update({
        "status": "Resolved",
        "resolution_notes": notes,
        "resolved_at": now_iso(),
        "resolved_by": data.get("resolved_by", "Admin"),
    }).eq("id", ticket_id).execute()

    log_activity(data.get("resolved_by", "Admin"), "resolved_incident", ticket.data[0]["ticket_id"], notes)
    return jsonify({"message": "Ticket resolved"}), 200


# --------------------------------------------------------------------------
# ACTIVITY LOG  ->  table: activity_log  (drives Recent Activity Feed)
# --------------------------------------------------------------------------
@app.route("/api/activity", methods=["GET"])
@require_supabase
def list_activity():
    limit = int(request.args.get("limit", 10))
    res = supabase.table("activity_log").select("*").order("created_at", desc=True).limit(limit).execute()
    return jsonify(res.data), 200


# --------------------------------------------------------------------------
# DASHBOARD STATS  ->  computed from documents / approvals / tickets / ai logs
# --------------------------------------------------------------------------
@app.route("/api/stats/overview", methods=["GET"])
@require_supabase
def stats_overview():
    docs = supabase.table("documents").select("id,status").execute().data or []
    approvals = supabase.table("document_approvals").select("id").eq("status", "Pending").execute().data or []
    tickets = supabase.table("maintenance_tickets").select("id,status,severity,created_at").execute().data or []
    ai_logs = supabase.table("ai_query_logs").select("sources,created_at").order("created_at", desc=True).limit(200).execute().data or []

    active_documents = len([d for d in docs if d.get("status") == "Active"])
    pending_approvals = len(approvals)
    open_incidents = len([t for t in tickets if t.get("status") != "Resolved"])
    critical_open = len([t for t in tickets if t.get("status") != "Resolved" and "P1" in (t.get("severity") or "")])

    today = datetime.datetime.utcnow().date()
    days = [today - datetime.timedelta(days=i) for i in range(6, -1, -1)]
    trend = []
    for d in days:
        count = 0
        for t in tickets:
            try:
                created = datetime.datetime.fromisoformat(t["created_at"]).date()
            except Exception:
                continue
            if created == d:
                count += 1
        trend.append({"day": d.strftime("%a")[0], "count": count})

    counter = collections.Counter()
    for log in ai_logs:
        for s in (log.get("sources") or "").split(", "):
            s = s.strip()
            if s:
                counter[s] += 1
    total = sum(counter.values()) or 1
    top_queries = [
        {"name": name, "pct": round(count / total * 100)}
        for name, count in counter.most_common(3)
    ]

    return jsonify({
        "active_documents": active_documents,
        "pending_approvals": pending_approvals,
        "open_incidents": open_incidents,
        "system_alerts": critical_open,
        "incident_trend": trend,
        "top_queries": top_queries,
    }), 200


# --------------------------------------------------------------------------
# PLANTMIND AI  ->  Groq, grounded on real extracted document text
# --------------------------------------------------------------------------
MODE_INSTRUCTIONS = {
    "general":    "Answer the question clearly and concisely for an employee using the company knowledge base.",
    "root_cause": "Perform a root cause analysis. Identify the most likely technical cause(s) of the issue described.",
    "policy":     "Answer strictly from a compliance/policy point of view, quoting relevant rules where applicable.",
    "summary":    "Give a short executive summary in 2-4 bullet points.",
}


def find_relevant_documents(question, limit=3):
    if supabase is None:
        return []
    try:
        res = supabase.table("documents") \
            .select("id,name,content,summary,version,file_path,file_type,file_name") \
            .eq("status", "Active") \
            .execute()
        docs = res.data or []
    except Exception:
        return []

    q_words = set(w.lower() for w in question.split() if len(w) > 2)

    scored = []
    for doc in docs:
        text = f"{doc.get('name','')} {doc.get('summary','')} {doc.get('content','')}".lower()
        score = sum(1 for w in q_words if w in text)
        if score == 0 and not (doc.get("content") or "").strip() and doc.get("file_path"):
            name_summary = f"{doc.get('name','')} {doc.get('summary','')}".lower()
            score = sum(1 for w in q_words if w in name_summary)
        if score > 0:
            scored.append((score, doc))

    scored.sort(key=lambda x: x[0], reverse=True)
    return [d for _, d in scored[:limit]]


@app.route("/api/ai/ask", methods=["POST"])
def ai_ask():
    if groq_client is None:
        return jsonify({"error": "Groq is not configured. Add GROQ_API_KEY to .env"}), 500

    data     = request.get_json(force=True) or {}
    question = (data.get("question") or "").strip()
    mode     = data.get("mode", "general")
    asked_by = data.get("asked_by", "Unknown")
    if not question:
        return jsonify({"error": "question is required"}), 400

    relevant_docs = find_relevant_documents(question)

    for d in relevant_docs:
        d["content"] = ensure_doc_content(d)

    context_text = "\n\n".join(
        f"[Document: {d['name']} ({d.get('version','v1.0')})]\n{(d.get('content') or d.get('summary') or '')[:6000]}"
        for d in relevant_docs
    ) or "No matching internal documents were found."

    instruction = MODE_INSTRUCTIONS.get(mode, MODE_INSTRUCTIONS["general"])

    system_prompt = f"""You are PlantMind AI, the internal IT knowledge assistant for OnlyTech.
{instruction}

Use ONLY the context below when it is relevant. If the context does not cover the question,
say so plainly and give general best-practice guidance instead.

--- CONTEXT ---
{context_text}
--- END CONTEXT ---

Respond in clear, well-formatted plain text (short paragraphs / bullet points where helpful)."""

    try:
        chat_completion = groq_client.chat.completions.create(
            model=GROQ_MODEL_NAME,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user",   "content": question},
            ],
            temperature=0.4,
            max_tokens=1024,
        )
        answer_text = chat_completion.choices[0].message.content
    except Exception as e:
        return jsonify({"error": f"Groq request failed: {str(e)}"}), 502

    sources = [{"name": d["name"], "version": d.get("version", "v1.0")} for d in relevant_docs]

    if supabase is not None:
        try:
            supabase.table("ai_query_logs").insert({
                "id":        str(uuid.uuid4()),
                "question":  question,
                "mode":      mode,
                "answer":    answer_text,
                "sources":   ", ".join(s["name"] for s in sources),
                "asked_by":  asked_by,
                "created_at": now_iso(),
            }).execute()
        except Exception:
            pass

    return jsonify({
        "answer":  answer_text,
        "mode":    mode,
        "sources": sources,
        "model":   GROQ_MODEL_NAME,
    }), 200


# --------------------------------------------------------------------------
if __name__ == "__main__":
    port = int(os.getenv("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=True)